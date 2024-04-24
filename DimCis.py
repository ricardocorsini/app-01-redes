import docx
import numpy as np
import math
import pandas as pd

class Beam:
    
    beamsCreated = []
    
    def __init__(self, name, bw, d, Vk, gama_c, gama_c2, fywk, gama_s, fck, stirrupLeg):
        self.name = name
        self.bw = bw
        self.d = d
        self.Vk = Vk
        self.gama_c = gama_c    #minoração do fck
        self.gama_c2 = gama_c2  #majoração do Vk
        self.fywk = fywk
        self.gama_s = gama_s
        self.fck = fck
        self.stirrupLeg = stirrupLeg

        self.h = self.d + 5
        self.Vd = (self.Vk) * (self.gama_c2)
        
        Beam.beamsCreated.append(self)

    def __str__(self):
        return f'{self.name} - {self.bw} X {self.h}'

    def concrete_properties(self):
        fcd = self.fck / self.gama_c
        fywd = self.fywk / self.gama_s
        fctm = 0.3 * pow(self.fck, 2 / 3)
        fctk_inf = 0.7 * fctm
        fctd = fctk_inf / self.gama_c

        dictExport = {
            'fcd': fcd,
            'fywd': fywd,
            'fctm': fctm,
            'fctk_inf': fctk_inf,
            'fctd': fctd
        }

        return dictExport

    def compressed_cis(self):

        fcd = Beam.concrete_properties(self)['fcd']

        alphaV2 = 1 - (self.fck / 250)
        vrd2 = 0.27 * alphaV2 * (fcd / 10) * self.bw * self.d

        if self.Vd <= vrd2:
            statusCompressed = 'ok'
        else:
            statusCompressed = 'disapproved'

        dictExport = {
            'status': statusCompressed,
            'alphaV2': alphaV2,
            'vrd2': vrd2
        }

        return dictExport

    def tension_cis(self):
        fctd = Beam.concrete_properties(self)['fctd']
        fctm = Beam.concrete_properties(self)['fctm']
        fywd = Beam.concrete_properties(self)['fywd']

        Vc = 0.6 * (fctd / 10) * self.bw * self.d

        #Reinforced min

        Asw_cm_min = 0.2 * (fctm) * (self.bw / self.fywk)
        Asw_m_min = Asw_cm_min * 100
        Vsw_min = Asw_cm_min * 0.9 * self.d * (fywd / 10)
        vrd3_min = Vsw_min + Vc

        if self.Vd <= vrd3_min:
            Asw_cis = Asw_m_min
            Asw_m = 0
            Vsw = 0
        else: 
            Vsw = self.Vd - Vc
            Asw_cm = (Vsw) / (0.9 * self.d * fywd / 10)
            Asw_m = Asw_cm * 100
            Asw_cis = Asw_m


        dictExport = {
            'Vc': Vc,
            'Vd': self.Vd,
            'Asw_cm_min': Asw_cm_min,
            'Asw_m_min': Asw_m_min,
            'Vsw_min': Vsw_min,
            'Vrd3_min': vrd3_min,
            'Vsw': Vsw,
            'Asw_m': Asw_m,
            'Asw_cis': Asw_cis,
        }

        return dictExport
    
    def debiting(self):

        Asw_cis = Beam.tension_cis(self)['Asw_cis']

        listDiameter = np.array([5.0, 6.3, 8.0, 10.0, 12.5, 16.0, 20.0, 25.0, 32.0, 40.0])
        areas = self.stirrupLeg * np.pi * pow((listDiameter / 10), 2) / 4 

        s = (areas / Asw_cis) * 100
        s_round = np.floor(s)
    

        dictExport = {
            'Diameter': listDiameter,
            'Spacing': s_round
        }

        dataFrame = pd.DataFrame(dictExport)

        return dataFrame


    def listar_vigas():
        for beam in Beam.beamsCreated:
            print(beam)

    def generate_memory(self):

        doc = docx.Document()
        doc.add_heading(f'{self.name} - {self.bw} x {self.h} - Cisalhamento', 0)

        p = doc.add_paragraph()
        p.add_run('Dimensionamento realizado pelo Modelo I - NBR 6118/2023').bold = True

        p = doc.add_paragraph()
        p.add_run('a) Dados de Entrada').bold = True

        for key, value in vars(self).items():
            if type(value) == float:
                value_round = round(value, 2)
            else:
                value_round = value

            doc.add_paragraph(f'{key} = {value_round}')

        doc.add_paragraph('')

        p = doc.add_paragraph()
        p.add_run('b) Dados do Concreto').bold = True

        for key, value in Beam.concrete_properties(self).items():
            if type(value) == float:
                value_round = round(value, 2)
            else:
                value_round = value

            doc.add_paragraph(f'{key} = {value_round}')

        doc.add_paragraph('')
        p = doc.add_paragraph()
        p.add_run('c) Bielas Comprimidas').bold = True

        for key, value in Beam.compressed_cis(self).items():
            doc.add_paragraph(f'{key} = {value}')

        doc.add_paragraph('')
        p = doc.add_paragraph()
        p.add_run('d) Bielas Tracionadas').bold = True

        for key, value in Beam.tension_cis(self).items():
            if type(value) == float:
                value_round = round(value, 2)
            else:
                value_round = value

            doc.add_paragraph(f'{key} = {value_round}')

        doc.add_paragraph('')
        p = doc.add_paragraph()
        p.add_run('e) Escolha de Armaduras').bold = True

        #Tabela

        table = doc.add_table(rows=(len(Beam.debiting(self)) + 1), cols=len(Beam.debiting(self).columns))
        table.style = 'Table Grid'
        for i, coluna in enumerate(Beam.debiting(self).columns):
            table.cell(0, i).text = coluna
        for i, row in enumerate(Beam.debiting(self).values):
            for j, value in enumerate(row):
                table.cell(i + 1, j).text = str(value)


        nameDoc = self.name + '_' + 'cis' + '.docx'
        doc.save(nameDoc)
